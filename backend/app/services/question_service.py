import re
from typing import Dict, List, Any, Optional
from pathlib import Path

from docx import Document
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker
from fastapi import HTTPException

from ..core.config import settings
from ..models.question import Question
from ..db.database import SessionLocal


class DocumentParser:
    """Parser for .docx files containing questions"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise HTTPException(status_code=400, detail="File not found")
        
        try:
            self.doc = Document(file_path)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Cannot read document: {str(e)}")
    
    def extract_metadata(self) -> Dict[str, str]:
        """Extract subject and lecturer from document paragraphs"""
        metadata = {"subject": "", "lecturer": ""}
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text.startswith("Subject:") and len(text.split()) > 1:
                metadata["subject"] = text.split()[1]
            elif text.startswith("Lecturer:") and len(text.split()) > 1:
                metadata["lecturer"] = text.split()[1]
        
        return metadata
    
    def parse_content_with_image(self, content: str) -> Dict[str, str]:
        """Parse content that may contain image references"""
        result = {"content": "", "content_img": ""}
        
        image_match = re.search(r"\[file:(.+?)\]", content)
        if image_match:
            result["content_img"] = image_match.group(1)
            result["content"] = re.sub(r"\[file:.*?\]", "", content).strip()
        else:
            result["content"] = content
        
        return result
    
    def parse_question_row(self, row_text: List[str], current_q: Dict[str, Any]) -> Dict[str, Any]:
        """Parse a single row and update question data"""
        if len(row_text) < 2:
            return current_q
            
        cell_text = row_text[0]
        value = row_text[1]
        
        if cell_text.startswith("QN="):
            current_q["code"] = row_text[0]
            content_data = self.parse_content_with_image(value)
            current_q.update(content_data)
            
        elif cell_text.startswith("ANSWER:"):
            current_q["answer"] = value
            
        elif cell_text.startswith("MARK:"):
            try:
                current_q["mark"] = float(value)
            except ValueError:
                current_q["mark"] = 0.0
                
        elif cell_text.startswith("UNIT:"):
            current_q["unit"] = value
            
        elif cell_text.startswith("MIX CHOICES:"):
            current_q["mix"] = value.lower() == 'yes'
            
        elif cell_text.startswith("a."):
            current_q["choiceA"] = value
        elif cell_text.startswith("b."):
            current_q["choiceB"] = value
        elif cell_text.startswith("c."):
            current_q["choiceC"] = value
        elif cell_text.startswith("d."):
            current_q["choiceD"] = value
            
        return current_q
    
    def parse_questions(self) -> List[Dict[str, Any]]:
        """Parse all questions from document tables"""
        metadata = self.extract_metadata()
        questions = []
        
        for table in self.doc.tables:
            current_q = {}
            
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not row_text:
                    continue
                
                # Start new question
                if row_text[0].startswith("QN="):
                    if current_q:  # Save previous question
                        current_q.update(metadata)
                        questions.append(current_q)
                    current_q = {}
                
                current_q = self.parse_question_row(row_text, current_q)
            
            # Add the last question
            if current_q:
                current_q.update(metadata)
                questions.append(current_q)
        
        return questions


class QuestionService:
    """Service for question operations"""
    
    @staticmethod
    def reading_file(file_path: str) -> List[Dict[str, Any]]:
        """Read and parse questions from .docx file"""
        parser = DocumentParser(file_path)
        return parser.parse_questions()
    
    @staticmethod
    def create_question_from_dict(question_data: Dict[str, Any]) -> Question:
        """Create Question model instance from dictionary"""
        return Question(
            code=question_data.get('code', ''),
            content=question_data.get('content', ''),
            content_img=question_data.get('content_img', ''),
            choiceA=question_data.get('choiceA', ''),
            choiceB=question_data.get('choiceB', ''),
            choiceC=question_data.get('choiceC', ''),
            choiceD=question_data.get('choiceD', ''),
            answer=question_data.get('answer', ''),
            mark=question_data.get('mark', 0.0),
            unit=question_data.get('unit', ''),
            subject=question_data.get('subject', ''),
            lecturer=question_data.get('lecturer', ''),
            importer=question_data.get('importer', 0),
            mix=question_data.get('mix', False)
        )

    @staticmethod
    def import_data(list_data: List[Dict[str, Any]], db: Optional[Session] = None) -> Dict[str, Any]:
        """Save questions to database"""
        if not list_data:
            return {"code": 400, "message": "No data to import"}
        
        # Use provided session or create new one
        session = db or SessionLocal()
        should_close = db is None
        
        try:
            questions = [QuestionService.create_question_from_dict(item) for item in list_data]
            session.add_all(questions)
            session.commit()
            
            return {
                "code": 200, 
                "message": f"Successfully imported {len(questions)} questions"
            }
            
        except Exception as e:
            session.rollback()
            return {
                "code": 500, 
                "message": f"Import failed: {str(e)}"
            }
        finally:
            if should_close:
                session.close()

    @staticmethod
    def get_questions(skip: int = 0, limit: int = 100, db: Optional[Session] = None) -> List[Question]:
        """Get questions from database with pagination"""
        session = db or SessionLocal()
        should_close = db is None
        
        try:
            stmt = select(Question).offset(skip).limit(limit)
            return session.execute(stmt).scalars().all()
        finally:
            if should_close:
                session.close()


# Backward compatibility functions
def reading_file(file_path: str) -> List[Dict[str, Any]]:
    return QuestionService.reading_file(file_path)

def import_data(list_data: List[Dict[str, Any]], db: Optional[Session] = None) -> Dict[str, Any]:
    return QuestionService.import_data(list_data, db)

def get_question(skip: int = 0, limit: int = 100):
    return QuestionService.get_questions(skip, limit)
