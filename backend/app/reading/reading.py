import re

from docx import Document
from fastapi.responses import JSONResponse
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker

from ..core.config import settings
from ..models.question import Question


# read data from .docx file with provided file path
def reading_file(file_path):
    doc = Document(file_path)
    subject = ""
    lecturer = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Subject:"):
            subject = text.split()[1]
        if text.startswith("Lecturer:"):
            lecturer = text.split()[1]

    questions = []
    for table_idx, table in enumerate(doc.tables):
        current_q = {}

        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not row_text:
                continue  # ignore empty row

            cell_text = row_text[0]

            # Check data and put it into a dict
            if cell_text.startswith("QN="):
                # Save the old question
                if current_q:
                    questions.append(current_q)
                    current_q = {}

                current_q["code"] = row_text[0]

                if re.search(r"\[file:(.+?)\]", row_text[1]):
                    match = re.search(r"\[file:(.+?)\]", row_text[1])
                    if match:
                        current_q["content_img"] = match.group(1)
                    # Get text before [file:...]
                    before_file = re.sub(r"\[file:.*?\]", "", row_text[1]).strip()
                    if before_file:
                        if "question" not in current_q:
                            current_q["content"] = before_file
                        else:
                            current_q["content"] += " " + before_file
                else:
                    current_q["content"] = row_text[1]
                    current_q["content_img"] = ""

            elif cell_text.startswith("ANSWER:"):
                current_q["answer"] = row_text[1]

            elif cell_text.startswith("MARK:"):
                try:
                    current_q["mark"] = float(row_text[1])
                except ValueError:
                    current_q["mark"] = 0

            elif cell_text.startswith("UNIT:"):
                current_q["unit"] = row_text[1]

            elif cell_text.startswith("MIX CHOICES:"):
                if row_text[1] == "Yes":
                    current_q["mix"] = True
                else:
                    current_q["mix"] = False

            elif cell_text.startswith("a."):
                current_q["choiceA"] = row_text[1]

            elif cell_text.startswith("b."):
                current_q["choiceB"] = row_text[1]

            elif cell_text.startswith("c."):
                current_q["choiceC"] = row_text[1]

            elif cell_text.startswith("d."):
                current_q["choiceD"] = row_text[1]

            current_q["subject"] = subject
            current_q["lecturer"] = lecturer

        # Add the last question
        if current_q:
            questions.append(current_q)

    return questions


# Todo: save data to the database
def import_data(list_data):
    # Create session
    engine = create_engine(settings.DATABASE_URL)
    Session = sessionmaker(bind=engine)
    session = Session()

    questions = []
    # bind data
    for item in list_data:
        question = Question(
            code=item["code"],
            content=item["content"],
            content_img=item["content_img"],
            choiceA=item["choiceA"],
            choiceB=item["choiceB"],
            choiceC=item["choiceC"],
            choiceD=item["choiceD"],
            answer=item["answer"],
            mark=item["mark"],
            unit=item["unit"],
            subject=item["subject"],
            lecturer=item["lecturer"],
            mix=item["mix"],
        )
        questions.append(question)

    # add item into database
    try:
        session.add_all(questions)
        session.commit()
        print("Import data successfully")
        return {"code": 200, "message": "Import thành công!"}
    except Exception as e:
        session.rollback()
        print("Có lỗi xảy ra:", str(e))
        return {"code": 202, "message": "Lỗi xảy ra " + str(e)}
    finally:
        session.close()


def get_question(skip: int = 0, limit: int = 100):
    # Create session
    engine = create_engine(settings.DATABASE_URL)
    Session = sessionmaker(bind=engine)
    session = Session()
    stmt = select(Question)
    stmt = stmt.offset(skip).limit(limit)
    return session.execute(stmt).scalars().all()
